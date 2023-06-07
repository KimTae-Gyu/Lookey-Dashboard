from flask import Flask, jsonify, request, send_file
import requests
import io
from datetime import datetime
from docx import Document
from flask_cors import CORS
import base64

# ----------- 보고서 생성-------------
exportDoc = Document('pic\알람 조치 내역 보고서.docx')

# 문서 안의 모든 표 가져오기
table = exportDoc.tables

# table[0]는 담당자 사인 등 문서 가장 위의 내용 
# table[1]은 순서대로 알람시각, 조치자, 조치시간, 조치여부
alarmTime = table[1].rows[1].cells[0].paragraphs[0]
checkPerson = table[1].rows[1].cells[1].paragraphs[0]
checkTime = table[1].rows[1].cells[2].paragraphs[0]
checkWhether = table[1].rows[1].cells[3].paragraphs[0]

# table[2]는 순서대로 알람시각, 알람내용, 위치, 원인 IP, 시도 행위, 위험도
eventTime = table[2].rows[0].cells[1].paragraphs[0]
alarmDesc = table[2].rows[1].cells[1].paragraphs[0]
alarmLoc = table[2].rows[2].cells[1].paragraphs[0]
alarmIp = table[2].rows[3].cells[1].paragraphs[0]
trialAction = table[2].rows[4].cells[1].paragraphs[0]
severity = table[2].rows[5].cells[1].paragraphs[0]

# table[3]는 순서대로 조치시각, 차단 IP
actionTime = table[3].rows[0].cells[1].paragraphs[0]
blockIp = table[3].rows[1].cells[1].paragraphs[0]
                                               
# table[4]는 비고
note = table[4].rows[0].cells[1].paragraphs[0]

def testDocs(inputData):
    alarmTime.text = 'AM11:00'
    checkPerson.text = '김지원'
    checkTime.text = 'AM11:01'
    checkWhether.text = '확인'

    eventTime.text = 'AM11:00'
    alarmDesc.text = '이상행위 탐지'
    alarmLoc.text = 'WAF'
    alarmIp.text = '123.234.213.111'
    trialAction.text = 'RuleSet : coreRule'
    severity.text = 'high'

    actionTime.text = 'AM11:01'
    blockIp.text = '123.234.213.111'

    note.text = inputData
    print("inputData=", inputData)

    # 현재 시각을 기반으로 파일 이름 생성
    currentTime = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Report_{currentTime}.docx"
    print("currentTime=", currentTime)
    print("filename=", filename)

    memoryFile = io.BytesIO()
    print("memoryFile=", memoryFile)
    exportDoc.save(memoryFile)
    memoryFile.seek(0)

    sendFile = memoryFile.getvalue()
    # print("sendFile=", sendFile)
    return sendFile

# ---------------------서버 설정-----------------

app = Flask(__name__)
CORS(app)

@app.route('/sendData', methods=['POST'])
def sendData():
    # Word 문서 데이터 가져오기
    jsonInputData = request.json.get('noteWrite', '')
    # print("jsonInputData:",jsonInputData)
    sendFile = testDocs(jsonInputData)
    # print("sendFile:", sendFile)
    
    # Node.js 서버로 데이터 전송
    nodejsUrl = "http://localhost:3000/processData"  # Node.js 서버 URL

    # 파일을 분할하여 여러 개의 청크로 나누기
    CHUNK_SIZE = 256 * 1024  # 256KB
    chunks = [sendFile[i:i + CHUNK_SIZE] for i in range(0, len(sendFile), CHUNK_SIZE)]
    # print("chunks:", chunks)
    chunksLen = len(chunks)
    # print("chunksLen:", chunksLen)

    base64Data = base64.b64encode(chunks)
    for index, chunk in enumerate(base64Data, start=1):
        # 바이너리 데이터 변경
        print("index:",index)
        # print("chunk:", chunk)
        # print("base64Data:",base64Data)
        # 분할된 데이터와 함께 요청 전송
        response = requests.post(nodejsUrl, json={'chunkData': base64Data.decode('utf-8'), 'chunk':chunk, 'chunkIndex': index, 'chunksLen': chunksLen}, headers={'Content-Type': 'application/json'})
        print("response(node.js Url):", response)

        if response.status_code != 200:
            return jsonify({'success': False, 'error': 'Failed to communicate with Node.js server'})

    # 반복문 종료 후에도 response 변수에 접근 가능
    if response.status_code == 200:
        # Node.js 서버에서 전달받은 응답 데이터 확인
        downloadLink = response.json().get('downloadLink')
        print("downloadLink:", downloadLink)
        downloadUrl = f"/download?file={downloadLink}"
        print("downloadUrl:", downloadUrl)
        if downloadLink:
            return jsonify({'success': True, 'downloadLink': downloadUrl})
        else:
            return jsonify({'success': False, 'error': 'Failed to generate download link'})
    else:
        return jsonify({'success': False, 'error': 'Failed to communicate with Node.js server'})

if __name__ == '__main__':
    app.run()
