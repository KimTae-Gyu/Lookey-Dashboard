from docx import Document
from flask import Flask
from flask import request
# from docx.shared import RGBColor

app=Flask(__name__)

exportDoc = Document('pic\알람 조치 내역 보고서.docx')

# 문서 안의 모든 표 가져오기
table = exportDoc.tables

# table[0]는 담당자 사인 등 문서 가장 위의 내용 
# table[1]은 순서대로 알람시각, 조치자, 조치시간, 조치여부
alarmTime = table[1].rows[1].cells[0].paragraphs[0]
checkPerson = table[1].rows[1].cells[1].paragraphs[0]
checkTime = table[1].rows[1].cells[2].paragraphs[0]
checkWhether = table[1].rows[1].cells[3].paragraphs[0]

# table[2]는 순서대로 알람시각, 알람내용, 위치, 원인 IP, 시도 행위, 위험도
eventTime = table[2].rows[0].cells[1].paragraphs[0]
alarmDesc = table[2].rows[1].cells[1].paragraphs[0]
alarmLoc = table[2].rows[2].cells[1].paragraphs[0]
alarmIp = table[2].rows[3].cells[1].paragraphs[0]
trialAction = table[2].rows[4].cells[1].paragraphs[0]
severity = table[2].rows[5].cells[1].paragraphs[0]

# table[3]는 순서대로 조치시각, 차단 IP
actionTime = table[3].rows[0].cells[1].paragraphs[0]
blockIp = table[3].rows[1].cells[1].paragraphs[0]
                                               

# table[4]는 비고
note = table[4].rows[0].cells[1].paragraphs[0]

@app.route("/", methods=["GET"])

def testDocs():
    # node_server_url= "http://127.0.0.1:3000/report"

    # try:
    #     response = requests.get(node_server_url)  # Node.js 서버에 GET 요청 보내기
    #     response_data = response.json()  # JSON 형식의 응답 데이터 가져오기

    #     # 응답 데이터 처리
    #     # 예를 들어, 응답 데이터의 필드에 접근하여 값을 가져올 수 있습니다.
    #     data_value = response_data["data_field"]


    alarmTime.text = response_data["data_field"]
    checkPerson.text = response_data["data_field"]
    checkTime.text = response_data["data_field"]
    checkWhether.text = '확인'

    eventTime.text = response_data["data_field"]
    alarmDesc.text = '이상행위 탐지'
    alarmLoc.text = 'WAF'
    alarmIp.text = response_data["data_field"]
    trialAction.text = response_data["data_field"]
    severity.text = 'high'

    actionTime.text = response_data["data_field"]
    blockIp.text = response_data["data_field"]

    # d5_export.js의 텍스트 데이터를 가져옴
    a = input('텍스트를 입력하세요')
    note.text = a

    exportDoc.save('테스트보고서.docx')

    # alarmTime.text = 'AM11:00'
    # checkPerson.text = '김지원'
    # checkTime.text = 'AM11:01'
    # checkWhether.text = '확인'

    # eventTime.text = 'AM11:00'
    # alarmDesc.text = '이상행위 탐지'
    # alarmLoc.text = 'WAF'
    # alarmIp.text = '123.234.213.111'
    # trialAction.text = 'RuleSet : coreRule'
    # severity.text = 'high'

    # actionTime.text = 'AM11:01'
    # blockIp.text = '123.234.213.111'

    # a = input('텍스트를 입력하세요')
    # note.text = a

    # exportDoc.save('테스트보고서.docx')

if __name__ == "__main__":
    app.run()


# exportDoc.save('알람 조치 내역 보고서.docx')